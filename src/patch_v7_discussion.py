#!/usr/bin/env python3
"""Patch v7: add Discussion bridge paragraph on drug-level heterogeneity."""
import json
from docx import Document

SRC = "/data/data/Drug_Pred/research/PathOmicDRP_Full_Manuscript_v7.docx"
DISC = (
    "The sharp drug-level split we observed in the fair PCA-256 comparison is, we "
    "believe, one of the more interesting practical observations of this study. It "
    "reframes the question from 'does multi-modal fusion improve drug-response "
    "prediction?' to 'for which drugs does multi-modal fusion improve drug-response "
    "prediction?' and supplies a testable mechanistic hypothesis — fusion wins when "
    "response determinants span modalities — that can be revisited in larger "
    "prospective cohorts. Because different datasets sample different drugs with "
    "different response-determinant structures, the apparently contradictory "
    "direction of embedding-versus-baseline advantage across drug panels "
    "(e.g. fusion superior on doxorubicin and cyclophosphamide but inferior on "
    "taxanes) is not a failure of reproducibility but an expected drug-specific "
    "phenomenon. We therefore position PathOmicDRP not as a universal drug-response "
    "predictor but as a method whose clinical utility should be matched to the "
    "mechanistic class of the drug being evaluated."
)

def main():
    doc = Document(SRC)
    anchor_substr = "The critical test is whether models trained on imputed targets"
    for p in doc.paragraphs:
        if anchor_substr in p.text:
            anchor_el = p._p
            parent = anchor_el.getparent()
            idx = list(parent).index(anchor_el)
            new_p = doc.add_paragraph(DISC, style='Normal')
            parent.remove(new_p._p)
            parent.insert(idx + 1, new_p._p)
            break
    doc.save(SRC)
    print(f"Patched {SRC}")

if __name__ == '__main__':
    main()
