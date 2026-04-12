#!/usr/bin/env python3
"""Patch v7 with new table legends and figure legends."""
from docx import Document

SRC = "/data/data/Drug_Pred/research/PathOmicDRP_Full_Manuscript_v7.docx"


def add_paragraph_after(doc, anchor_substr, texts, style='Normal'):
    if isinstance(texts, str): texts = [texts]
    for p in doc.paragraphs:
        if anchor_substr in p.text:
            anchor_el = p._p
            parent = anchor_el.getparent()
            idx = list(parent).index(anchor_el)
            created = [doc.add_paragraph(t, style=style) for t in texts]
            for np_ in created: parent.remove(np_._p)
            for j, np_ in enumerate(created):
                parent.insert(idx + 1 + j, np_._p)
            return True
    return False


def replace_paragraph_text(doc, old_substring, new_text):
    for p in doc.paragraphs:
        if old_substring in p.text:
            for r in p.runs: r.text = ""
            if p.runs: p.runs[0].text = new_text
            else: p.add_run(new_text)
            return True
    return False


def main():
    doc = Document(SRC)

    # Update Table 3 legend: replace "best four-modal model" with CV-averaged framing
    replace_paragraph_text(
        doc,
        "Modality importance quantified by PCC drop upon ablation of each modality in the best",
        "Table 3. Modality importance quantified by per-modality PCC drop across 5-fold "
        "cross-validation (mean ± SD, 95% CI). Replaces an earlier single-fold estimate. "
        "Columns: mean ΔPCC_drug upon inference-time zeroing of each modality; 95% bootstrap "
        "confidence interval across folds; share of total summed ablation drop."
    )

    # Insert Table 4 and Table 5 legends after Table 3
    add_paragraph_after(
        doc,
        "share of total summed ablation drop",
        [
            "Table 4. Per-drug clinical AUC comparison (fair embedding). For each of the four "
            "drugs with sufficient TCGA treatment-outcome labels, the mean AUC (3-fold "
            "stratified CV) and bootstrap 95% CI are reported for five predictors: PathOmicDRP "
            "learned 256-dim embedding, PCA-256 of raw 4-modal features (dimensionality-matched "
            "baseline), PCA-256 of omics-only, raw 4-modal features, and oncoPredict-imputed "
            "IC₅₀ (13 drugs). Winner column: Δ>0.05 favors PathOmicDRP (★), Δ<−0.05 favors "
            "PCA-256 baseline. Source: drug_winner_table.csv.",

            "Table 5. Per-drug imputed-IC₅₀ histopathology contribution classified by "
            "mechanism of action (13 drugs, 5-fold CV mean). Columns: 4-modal PCC, 3-modal "
            "PCC, ΔPCC (histology contribution), and the number of CV folds in which histology "
            "improved prediction. Grouped by MOA class (DNA-damaging, Microtubule, Hormone, "
            "Targeted, Apoptosis). Source: imputed_ic50_landscape.csv.",
        ],
    )

    # Add Figure Legends for revised Fig 5A, and new Fig 9, S17, S18 (after existing S16 legend)
    add_paragraph_after(
        doc,
        "Figure S16. External validation against GDSC pharmacological data",
        [
            "Figure 5A (revised). CV-averaged modality importance. Horizontal bar plot of mean "
            "ΔPCC_drug when each modality is zeroed at inference time, averaged across five "
            "cross-validation folds. Error bars show 1 SD; annotated intervals are bootstrap "
            "95% confidence intervals across folds. Replaces the previous single-fold estimate "
            "that attributed 89.5% of modality importance to histopathology; the CV-averaged "
            "figure is 79.8% [95% CI 0.20, 0.28] for histology, 18.4% for transcriptomics, "
            "and <2% each for genomic and proteomic tokens. Source: "
            "figures_v6/Fig5A_cv_ablation.pdf.",

            "Figure 9. Drug-level heterogeneity in fusion benefit. (A) Clinical-AUC advantage "
            "per drug (PathOmicDRP 256-dim learned embedding minus PCA-256 of raw 4-modal "
            "features) for the four drugs with sufficient TCGA treatment-outcome labels. "
            "Bars are annotated with each drug's mechanism-of-action class. DNA-damaging "
            "agents (cyclophosphamide, doxorubicin) show large positive advantages for "
            "PathOmicDRP; microtubule-targeting taxanes (docetaxel, paclitaxel) show large "
            "negative advantages (PCA-256 wins). (B) Imputed-IC₅₀ histopathology "
            "contribution per drug across all 13 drugs, coloured by MOA class. Histopathology "
            "helps on the imputed-IC₅₀ task for hormone, taxane and vinca-alkaloid drugs; "
            "contributions are neutral or slightly negative for cell-line-imputed DNA-damaging "
            "agents and two BCL-family inhibitors. Source: figures_v6/Fig9_drug_heterogeneity.pdf.",

            "Figure S17. METABRIC external validation (n = 1,980). (A) ER-status "
            "stratification of predicted Tamoxifen and Fulvestrant IC₅₀ (Mann–Whitney p ≈ "
            "10⁻³⁰ for both). (B) Drug-drug correlation-structure conservation between TCGA "
            "predictions and METABRIC predictions across 13 drugs (Spearman ρ = 0.961 on the "
            "upper-triangular correlation pairs, p = 4.4×10⁻⁴⁴, 78 pairs). (C) Cox "
            "proportional-hazards per-SD HRs for each predicted IC₅₀ on overall survival "
            "(n = 1,980, 150 death events); red dots denote BH-FDR q < 0.05, orange q < 0.1. "
            "Source: figures_v6/FigS17_metabric_validation.pdf.",

            "Figure S18. Fair-embedding clinical AUC comparison. For each of the four drugs "
            "with TCGA clinical-outcome labels, bars compare PathOmicDRP's learned 256-dim "
            "embedding against PCA-256 of raw 4-modal features, PCA-256 of omics-only, raw "
            "4-modal features, and imputed IC₅₀ (13-dim). Error bars are bootstrap 95% CIs "
            "over 2,000 replicates of the pooled CV test-fold predictions. The dashed line "
            "marks chance (AUC = 0.5). Source: figures_v6/FigS18_fair_embedding.pdf.",
        ],
    )

    doc.save(SRC)
    print(f"Patched {SRC}")


if __name__ == '__main__':
    main()
